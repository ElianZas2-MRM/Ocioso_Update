# -*- coding: utf-8 -*-
"""Reconstruye la sección Mejoras completa → FINAL3.docx"""
from docx import Document
from docx.oxml.ns import qn

doc = Document('Arquitectura_Consolidada_ Osocio Form Automation_FINAL2.docx')
body = doc.element.body

# ── Encontrar y eliminar desde el H1 "Mejoras" hasta el final ─────────────────
def get_para_text(elem):
    return ''.join(r.text for r in elem.iter(qn('w:t')) if r.text)

def is_h1(elem):
    if elem.tag != qn('w:p'): return False
    pPr = elem.find(qn('w:pPr'))
    if pPr is None: return False
    pStyle = pPr.find(qn('w:pStyle'))
    if pStyle is None: return False
    return 'Heading1' in pStyle.get(qn('w:val'), '')

mejoras_idx = None
for i, elem in enumerate(list(body)):
    if is_h1(elem) and 'Mejoras' in get_para_text(elem):
        mejoras_idx = i
        break

all_elems = list(body)
for elem in all_elems[mejoras_idx:]:
    if elem.tag != qn('w:sectPr'):
        body.remove(elem)

print(f'Eliminados elementos desde índice {mejoras_idx}')

# ── Helpers ───────────────────────────────────────────────────────────────────
def h1(t):  doc.add_heading(t, level=1)
def h2(t):  doc.add_heading(t, level=2)
def h3(t):  doc.add_heading(t, level=3)
def p(t):   doc.add_paragraph(t)
def b(t):   doc.add_paragraph('• ' + t)

# ══════════════════════════════════════════════════════════════════════════════
h1('Mejoras y Nuevas Funcionalidades')
p('Esta sección documenta en detalle todas las mejoras y nuevas funcionalidades incorporadas al sistema Osocio Form Automation. Incluye tanto las características actuales del programa como aquellas que reemplazan soluciones anteriores que fueron descontinuadas.')

# ── 1 ─────────────────────────────────────────────────────────────────────────
h2('1. Arquitectura refactorizada — Base genérica por país')
p('Archivos principales: core/generic_country_base.py y core/country_configs.py')
p('Anteriormente el sistema tenía 9 archivos individuales (Formulario_Argentina_Base.py, Formulario_Bolivia_Base.py, etc.), uno por cada país. Cada archivo era una clase separada con la configuración de campos hardcodeada. Cualquier cambio general debía replicarse manualmente en los 9 archivos.')
p('Con la refactorización:')
b('Se eliminaron los 9 archivos Base individuales por país.')
b('Se creó generic_country_base.py: una sola clase GenericCountryBase que hereda de BaseFormFiller y carga la configuración de cualquier país de forma dinámica.')
b('Se creó country_configs.py: un diccionario centralizado COUNTRY_CONFIGS con la configuración declarativa de los 9 países. Cada entrada define el archivo Excel, el field_mapping (lista de campos con tipo, ID HTML y columna del Excel) y los required_fields.')
b('Para agregar un nuevo país ahora solo se necesita agregar una nueva entrada en el diccionario, sin duplicar lógica.')
b('El archivo forms/_runner_common.py actúa como fábrica: crea dinámicamente el runner de cualquier país leyendo su configuración de COUNTRY_CONFIGS.')
p('Impacto: el código de automatización central vive en un solo lugar (base_form_filler.py) y las diferencias por país son puramente declarativas (datos, no lógica).')

# ── 2 ─────────────────────────────────────────────────────────────────────────
h2('2. Soporte multi-navegador y multi-viewport dinámico')
p('Archivos: core/browser_manager.py y interface/main_interface.py')
p('Versiones anteriores tenían 54 archivos separados (9 países × 6 combinaciones de navegador/viewport): Formulario_Argentina_Chrome_Desktop.py, Formulario_Argentina_Chrome_Mobile.py, Formulario_Argentina_Firefox_Desktop.py, etc. Ese enfoque era muy difícil de mantener: cualquier corrección debía aplicarse a 54 archivos.')
p('La arquitectura actual centraliza esto en dos componentes:')
b('browser_manager.py: recibe los parámetros browser_type (chrome/firefox/edge), viewport (fullscreen o 600x738 px) y headless (True/False) y construye el WebDriver correcto con los argumentos apropiados para cada navegador.')
b('La interfaz gráfica (main_interface.py) expone un selector de navegador y viewport. El usuario elige desde la UI y se pasa la selección al motor, sin necesitar archivos separados.')
b('Navegadores soportados: Google Chrome, Mozilla Firefox, Microsoft Edge.')
b('Viewports soportados: pantalla completa (fullscreen) y vista móvil (600 × 738 px, para simular un teléfono).')
p('Resultado: 54 archivos estáticos fueron reemplazados por 1 módulo configurable. Los archivos _Mov_*.py quedaron obsoletos y fueron eliminados del proyecto.')

# ── 3 ─────────────────────────────────────────────────────────────────────────
h2('3. Llenado automático de formularios multipaso (auto-step)')
p('Archivo: core/base_form_filler.py — método fill_form_fields_auto_step()')
p('Muchos formularios no muestran todos los campos en una sola pantalla, sino que los distribuyen en pasos (Paso 1: datos personales, Paso 2: vehículo, Paso 3: confirmación). Versiones anteriores requerían archivos específicos _3_Pasos.py para manejar esta situación. Esos archivos fueron eliminados porque el nuevo método los hace completamente innecesarios.')
p('La implementación actual resuelve esto de forma completamente automática:')
b('fill_form_fields_auto_step() llena todos los campos visibles del paso actual.')
b('Después de llenar los campos, busca automáticamente botones de avance cuyo texto contenga "Siguiente", "Seguinte" (portugués), "Next" o "Continuar".')
b('Si encuentra el botón, lo presiona y espera a que cargue el nuevo paso, luego repite el proceso.')
b('Itera hasta un máximo de 15 veces (configurable), lo que cubre tanto formularios de 1 paso como de múltiples pasos sin ninguna configuración extra.')
b('Si no hay botón de siguiente, asume que está en el paso final y procede al envío del formulario.')
p('Este mecanismo reemplazó por completo los archivos _3_Pasos.py, que fueron eliminados del proyecto al volverse redundantes.')

# ── 4 ─────────────────────────────────────────────────────────────────────────
h2('4. Auto-relleno de dropdowns y campos no mapeados')
p('Archivo: core/base_form_filler.py — método _auto_fill_unmapped_dropdowns()')
p('A veces los formularios añaden campos nuevos que no están en el field_mapping configurado (porque el formulario cambió en producción y el mapping no se actualizó todavía). Sin esta mejora, el formulario quedaría incompleto y el envío fallaría por campos requeridos vacíos.')
p('El método _auto_fill_unmapped_dropdowns() actúa como red de seguridad:')
b('Recorre todos los elementos <select> e <input> visibles en el DOM actual.')
b('Para cada elemento que no está en el field_mapping, verifica si hay un valor definido en los IDs dinámicos (ids_dinamicos.json). Si lo hay, usa ese valor.')
b('Si no hay un valor dinámico configurado, para dropdowns selecciona automáticamente una opción válida aleatoria (excluyendo placeholders "Seleccione...", "Choose...", y opciones deshabilitadas o con valor vacío).')
b('Retorna True si completó al menos un campo, False si no había nada por completar.')
p('Esto hace que el sistema sea resiliente ante cambios parciales del formulario en producción: puede seguir funcionando aunque aparezcan campos nuevos no mapeados, y el operador puede actualizar el mapping en su próxima revisión sin urgencia.')

# ── 5 ─────────────────────────────────────────────────────────────────────────
h2('5. Tracking y escritura de campos dinámicos en el Excel de resultados')
p('Archivo: core/base_form_filler.py — método write_tracked_fields_to_sheet()')
p('El Excel de resultados no solo registra si el formulario fue enviado con éxito o no. También guarda qué valores se usaron realmente para cada campo durante la ejecución. Esto es fundamental para auditoría y debugging.')
p('Funcionamiento:')
b('A medida que el motor llena cada campo del formulario, registra internamente el par (campo, valor efectivamente usado) mediante _record_field_value().')
b('Al finalizar cada fila, write_tracked_fields_to_sheet() escribe todos esos valores en el Excel de resultados.')
b('Para los campos que están en el field_mapping, escribe el valor en la columna correspondiente.')
b('Para los campos que NO estaban en el mapping (campos dinámicos detectados y rellenados automáticamente), crea columnas nuevas con el formato ID::{field_id}, donde field_id es el identificador HTML del campo en el formulario.')
p('Ejemplo: si el formulario tiene un campo con id="dealer-code" que no estaba mapeado, el Excel de resultados tendrá una columna nueva "ID::dealer-code" con el valor que el sistema seleccionó. Esto permite ver exactamente qué pasó en cada ejecución sin revisar logs ni capturas de pantalla.')

# ── 6 ─────────────────────────────────────────────────────────────────────────
h2('6. requested_data_index — alineación correcta de columnas del Excel')
p('Archivo: core/country_configs.py + core/base_form_filler.py — método _build_effective_data_keys()')
p('Cada campo en el field_mapping tiene un data_index que indica en qué columna del Excel está su valor de entrada. El problema ocurría cuando el orden de las columnas en el Excel no coincidía con el orden del array field_mapping: el sistema leía el dato de la columna equivocada.')
p('La solución es el campo opcional requested_data_index en cada entrada del field_mapping:')
b('Si un campo tiene requested_data_index definido, ese valor tiene prioridad sobre data_index para la lectura del Excel, independientemente de dónde esté el campo en el array.')
b('Esto permite que el orden del array field_mapping sea completamente independiente del orden de las columnas en el Excel.')
b('Se puede reordenar el array (por motivos de dependencias o presentación) sin romper la lectura de datos del Excel.')
b('También permite que dos campos lean de la misma columna del Excel si la situación lo requiere.')
p('Este mecanismo fue la corrección central del commit "Fix: Actualizar mapeo de índices del Excel para todos los formularios", que resolvió un bug que afectaba a todos los países cuando se reorganizaban columnas en el Excel.')

# ── 7 ─────────────────────────────────────────────────────────────────────────
h2('7. Inyección JavaScript para campos con máscara de formato — Brasil')
p('Archivo: core/base_form_filler.py — lógica específica para campos de Brasil')
p('Los formularios brasileños usan máscaras de entrada (input masks) en los campos de CPF, CNPJ y teléfono. Estas máscaras interceptan el teclado virtual de Selenium y filtran los caracteres ingresados, causando que send_keys() inserte valores incorrectos, truncados o directamente vacíos.')
p('La solución usa inyección directa en el DOM:')
b('En lugar de usar send_keys(), el sistema inyecta el valor directamente con driver.execute_script(), asignando el valor limpio (solo dígitos, sin guiones ni puntos) a la propiedad value del elemento HTML.')
b('Primero limpia el campo (element.value = ""), luego asigna el nuevo valor (element.value = "12345678901").')
b('Luego dispara manualmente los eventos "input" y "change" usando dispatchEvent(), para que la máscara del formulario formatee el valor visualmente (agregando puntos y guiones) y para que el framework JavaScript del formulario (Angular, React, etc.) detecte el cambio de estado.')
b('El proceso simula exactamente lo que haría un usuario pegando texto con Ctrl+V en el campo, evitando la captura carácter por carácter que bloquean las máscaras.')
p('Sin esta inyección, los campos CPF/CNPJ del formulario brasileño quedarían vacíos o con formato incorrecto, y el formulario rechazaría el envío por error de validación.')
p('Importante: este mecanismo de inyección JS aplica tanto cuando el valor viene del Excel como cuando el documento fue generado en tiempo de ejecución por la API de 4devs.')

# ── 8 ─────────────────────────────────────────────────────────────────────────
h2('8. Sanitización del número de documento según tipo — Perú')
p('Archivo: core/base_form_filler.py — método _sanitize_peru_document()')
p('El formulario de Perú tiene un campo "tipo de documento" antes del campo del número de documento. Cada tipo tiene un formato diferente: el DNI peruano tiene exactamente 8 dígitos, un RUC tiene 11 dígitos, un carné de extranjería tiene otra estructura, etc.')
p('El sistema maneja esto automáticamente:')
b('Primero selecciona el tipo de documento que corresponde al valor del Excel.')
b('Luego ajusta el valor del campo de número de documento para que cumpla el formato esperado por ese tipo seleccionado (por ejemplo, trunca o rellena con ceros hasta los 8 dígitos para DNI, o 11 para RUC).')
b('Esto evita errores de validación en el formulario causados por incompatibilidad entre el tipo elegido y el formato del número ingresado.')
p('Sin esta sanitización, el formulario peruano mostraría un mensaje de validación indicando que el número no corresponde al tipo seleccionado, y el envío fallaría.')

# ── 9 ─────────────────────────────────────────────────────────────────────────
h2('9. Detección de envío exitoso y TY page')
p('Archivo: core/base_form_filler.py — métodos submit_and_verify_form() y _resolve_submit_button()')
p('Después de completar todos los campos del formulario, el sistema necesita enviarlo y determinar si el envío fue exitoso. Esta detección es crítica: el resultado de cada fila del Excel depende de ella.')
p('Localización del botón de envío:')
b('_resolve_submit_button() busca el botón de envío usando múltiples selectores CSS (button[type="submit"], input[type="submit"], botones con texto "Enviar", "Submit", "Envio", etc.).')
b('Si los selectores CSS no encuentran nada, cae a una búsqueda por XPath genérica como fallback.')
b('Una vez localizado, hace scroll hasta él y lo presiona con JavaScript para evitar problemas de visibilidad.')
p('Sistema de verificación en dos niveles:')
b('Detección primaria — confirmación positiva: espera hasta 15 segundos a que aparezca el elemento div#thank-you con display: block en su estilo inline. Este div es el marcador de "formulario enviado correctamente" que usan los formularios de GM. Si aparece, el envío fue exitoso.')
b('Fallback — confirmación negativa (si el timeout de 15s expira sin detectar el div): el sistema busca evidencia de falla. Primero busca elementos visibles con clases o estilos de error (clases "error", color rojo, clase "invalid"). Si los encuentra, retorna el mensaje de error.')
b('Si no hay errores visibles, llama a _has_visible_required_field(): verifica si alguno de los campos del formulario sigue visible en el DOM. Si los campos siguen ahí, el formulario no se envió.')
b('Si no hay TY div, no hay errores visibles y no hay campos del formulario visibles, retorna "Enviado sin confirmación TY Page", lo que activa el mecanismo de retry automático (ver sección 18).')
p('Cuando la TY page se detecta exitosamente:')
b('El sistema sale del contexto del iframe (switch_to.default_content()) y vuelve al contexto principal de la página.')
b('Reposiciona la vista a la URL del formulario.')
b('Captura un screenshot completo de la página de confirmación, guardado como landing_typage_{N}.png en la carpeta de evidencias.')
p('Valores de retorno de submit_and_verify_form():')
b('"Lead enviado correctamente" + nombre del screenshot → envío exitoso con TY page.')
b('"Error visual detectado: {mensaje}" + None → el formulario mostró errores de validación.')
b('"Formulario sigue visible (TY Page no detectada)" + None → el formulario no se envió.')
b('"Enviado sin confirmación TY Page" + None → estado ambiguo, activa retry.')
p('El valor de retorno se escribe en la columna "Resultado" del Excel de salida junto con el nombre del screenshot de la TY page en la columna "TY Page".')

# ── 10 ────────────────────────────────────────────────────────────────────────
h2('10. Generador de datos de prueba por país')
p('Archivo: utils/data_generator.py — función principal generar_fila_datos(pais)')
p('Para hacer pruebas de los formularios se necesitan datos de leads realistas (nombres, documentos, teléfonos) que pasen las validaciones del formulario web. Crearlos manualmente es tedioso y propenso a errores de formato.')
p('El módulo data_generator.py genera automáticamente filas completas de datos válidos por país. La función generar_fila_datos(pais) devuelve un diccionario con todos los campos listos para insertar directamente en el Excel de entrada. La interfaz gráfica expone este generador en la pestaña "Generar Excels": el operador elige el país y la cantidad de filas, y el sistema puebla el Excel con un clic.')
p('Componentes del generador:')
b('Nombres y apellidos: seleccionados aleatoriamente de dos listas internas — aproximadamente 60 primeros nombres (masculinos y femeninos) y 60 apellidos en español. La función generar_nombre() elige 1 o 2 nombres al azar; generar_apellido() elige 1 o 2 apellidos.')
b('Emails: generados por generar_email(nombre, apellido, pais). El formato es {nombre_sin_acentos}{apellido}{separador}{abrev_pais}{NN}@mrm.com, donde el separador es "." o "_" elegido al azar y NN es un número de 2 dígitos. Ejemplo: "carlasanchez_ar07@mrm.com".')
b('Teléfonos: generados por generar_celular(pais) con prefijos y longitudes correctas por país. Argentina: código de área real (11, 351, 261...) + dígitos hasta 10 total. Brasil: DDD (11-98) + "9" + 8 dígitos = 11 total. Chile: "9" + 8 dígitos = 9 total. Ecuador y Paraguay: "09" + 8 dígitos. Uruguay: "09" + 7 dígitos.')
p('Generadores de documentos de identidad — detalle por país:')
p('RUT chileno — generar_rut_chile() en utils/data_generator.py:')
b('El RUT (Rol Único Tributario) chileno tiene entre 7 y 8 dígitos más un dígito verificador (DV) que puede ser un número del 0 al 9 o la letra K.')
b('El sistema genera un número aleatorio entre 5.000.000 y 25.000.000 (rango de RUTs de personas naturales activos).')
b('El dígito verificador se calcula con el algoritmo Módulo 11: se multiplican los dígitos del número de derecha a izquierda por los factores [2, 3, 4, 5, 6, 7] (que se repiten cíclicamente). Se suman todos los productos. El DV es 11 menos el resto de dividir esa suma por 11. Si el resultado es 11, el DV es "0"; si es 10, el DV es "K"; en cualquier otro caso es el número mismo.')
b('El resultado final se devuelve como cadena de texto sin puntos ni guión. También existe generar_rut_chile_con_k() que genera solo RUTs cuyo DV es K, para pruebas específicas del campo de documento en Chile.')
p('Cédula de identidad ecuatoriana — generar_ci_ecuador() en utils/data_generator.py:')
b('La cédula ecuatoriana tiene exactamente 10 dígitos. Los dos primeros representan la provincia de emisión (01 a 24), el tercero debe estar entre 0 y 6 (identifica el tipo de persona), los siguientes 6 son dígitos aleatorios, y el décimo es el dígito verificador.')
b('El dígito verificador se calcula con el algoritmo Módulo 10: se multiplican los primeros 9 dígitos de izquierda a derecha por los coeficientes alternos [2, 1, 2, 1, 2, 1, 2, 1, 2]. Si el resultado de la multiplicación es mayor o igual a 10, se le resta 9. Se suman todos los resultados. El DV es (10 - (suma % 10)) % 10.')
b('El sistema elige una provincia aleatoria entre 01 y 24, un tercer dígito entre 0 y 6, y genera los 6 dígitos del medio aleatoriamente. Luego calcula el DV y devuelve la cadena completa de 10 dígitos.')
p('CPF brasileño — generar_cpf_brasil() en utils/data_generator.py (solo para el generador de Excels):')
b('IMPORTANTE: esta función se usa únicamente para poblar el Excel de entrada desde la pestaña "Generar Excels". Durante la ejecución real del formulario, si el campo CPF está vacío en el Excel, el sistema usa la API de 4devs (ver más abajo), no esta función local.')
b('El CPF tiene 11 dígitos: 9 dígitos base más 2 dígitos verificadores. Los 9 primeros se generan aleatoriamente.')
b('Primer dígito verificador: se multiplican los 9 dígitos base por los pesos [10, 9, 8, 7, 6, 5, 4, 3, 2]. Se suma el total. Si el resto de dividir por 11 es menor que 2, el DV es 0; si no, el DV es 11 menos el resto.')
b('Segundo dígito verificador: mismo proceso con los 10 dígitos (9 base + primer DV) multiplicados por los pesos [11, 10, 9, 8, 7, 6, 5, 4, 3, 2]. Misma regla: si el resto < 2, DV = 0; si no, DV = 11 - resto.')
b('El resultado final son 11 dígitos sin puntuación ni guiones (solo números), que el formulario formatea visualmente con su máscara de entrada.')
p('API de 4devs para Brasil en tiempo de ejecución — core/base_form_filler.py:')
b('Cuando el campo CPF, CNPJ o CEP del Excel está vacío para una fila de Brasil, el sistema no falla. En cambio, genera el documento solicitando uno nuevo a la API pública de 4devs (https://www.4devs.com.br), que garantiza datos válidos.')
b('La API recibe una petición HTTP POST a https://www.4devs.com.br/ferramentas_online.php con el parámetro "acao": "gerar_cpf" (CPF sin puntuación), "gerar_cnpj" (CNPJ con puntuación XX.XXX.XXX/XXXX-XX), o "gerar_cep" (código postal de São Paulo).')
b('El sistema usa la librería estándar de Python (urllib.request) para la petición, sin dependencias externas adicionales. Timeout de 8 segundos.')
b('El valor recibido de la API se inyecta en el campo usando JavaScript (ver sección 7) y queda registrado en el Excel de resultados mediante el sistema de tracking de campos.')
p('Otros países (Argentina, Bolivia, Colombia, Paraguay, Perú, Uruguay):')
b('Para estos países, generar_documento(pais) devuelve un número entero aleatorio entre 1.000.000 y 99.999.999, sin validación de dígito verificador, ya que sus formularios no validan el formato del CI con algoritmo matemático.')

# ── 11 ────────────────────────────────────────────────────────────────────────
h2('11. Integración con LambdaTest para ejecución en macOS real')
p('Archivos: lambdatest_mac/lt_runner.py, lambdatest_mac/lt_controller.py, lambdatest_mac/lt_excel_reader.py')
p('Para validar que los formularios funcionan correctamente en macOS (donde no se puede correr el ejecutable .exe de Windows), el sistema se integra con LambdaTest, una plataforma de testing en la nube que provee navegadores reales corriendo en hardware macOS.')
p('Componentes principales:')
b('lt_runner.py: orquestador principal. Replica toda la lógica de Osocio (manejo de cookies GM, scroll dinámico, resolución de dependencias entre dropdowns, selección exacta en selects) pero sobre un WebDriver remoto de LambdaTest en lugar de un driver local.')
b('lt_controller.py: API pública que expone la función run(pais, platform, log_fn, stop_event), invocable desde la interfaz gráfica para lanzar sesiones en macOS sin salir de la aplicación.')
b('lt_excel_reader.py: lector del Excel de leads adaptado al contexto del runner remoto, con manejo de rutas relativas al entorno de ejecución.')
b('Credenciales: el username y access_key de LambdaTest se leen desde json/config_global.json. Si no están configurados ahí, el sistema busca en un archivo lambdatest_credentials.txt como alternativa.')
p('¿Cómo funciona? — flujo paso a paso:')
b('1. El operador configura el username y access_key de LambdaTest en json/config_global.json (o en lambdatest_credentials.txt).')
b('2. Desde la interfaz gráfica, el operador selecciona el país y presiona el botón de ejecución Mac/LambdaTest. Esto llama a lt_controller.run(pais, platform, log_fn, stop_event).')
b('3. lt_controller construye las "capabilities" del navegador remoto: qué navegador usar, qué versión, que el sistema operativo sea macOS, y el nombre de la sesión para identificarla en el dashboard de LambdaTest.')
b('4. Se crea un RemoteWebDriver apuntando a la URL del Hub de LambdaTest (hub.lambdatest.com/wd/hub) usando las credenciales configuradas. En este punto LambdaTest asigna una máquina macOS real con el navegador solicitado.')
b('5. lt_runner.py toma el control del RemoteWebDriver y replica el flujo completo de Osocio: navega a la landing page, detecta el iframe del formulario, llena los campos del Excel usando inyección JavaScript (para evitar pérdida de caracteres por latencia de red), avanza por los pasos del formulario, y envía el lead.')
b('6. Al finalizar cada fila, los resultados (éxito o error) se guardan en resultados_lambdatestmac/ con el mismo formato Excel que la ejecución local.')
b('7. La sesión LambdaTest se cierra limpiamente al finalizar todas las filas del Excel.')
p('Nota importante: actualmente la integración con LambdaTest no incluye captura de screenshots en macOS. Las evidencias visuales de las ejecuciones en Mac no se toman en esta versión. El texto se ingresa siempre vía JavaScript puro (execute_script + dispatchEvent) en lugar de send_keys, porque el entorno remoto tiene latencia de red y send_keys puede perder caracteres en conexiones lentas.')

# ── 12 ────────────────────────────────────────────────────────────────────────
h2('12. Módulo de validación automática de formularios (QA)')
p('Archivos: validation/selenium_validation_runner.py, validation/error_message_validator.py, validation/text_field_validator.py, validation/validation_exporter.py, validation/validation_email.py')
p('Independientemente de la automatización de carga de leads, el sistema incluye un subsistema completo de QA que verifica que los campos del formulario se comportan correctamente: que muestran los mensajes de error apropiados, que aceptan y rechazan los valores correctos, y que los dropdowns dependientes funcionan como se espera.')
p('Componentes del módulo:')
b('selenium_validation_runner.py: motor principal de QA. Abre el formulario real con Selenium, descubre todos los campos visibles incluyendo los de pasos intermedios, y para cada campo ejecuta casos de prueba: campo vacío, valores inválidos, valores fuera de límite (min/max length), caracteres especiales, formatos incorrectos.')
b('error_message_validator.py: lee los mensajes de error que aparecen en el formulario después de cada caso de prueba. Compara el mensaje obtenido con el esperado según las reglas configuradas en el JSON del país. Genera los inputs de prueba a partir de reglas declarativas como "required", "min_length", "max_length", "invalid_chars".')
b('text_field_validator.py: validaciones específicas para campos de texto como email, número de documento y teléfono.')
b('validation_email.py: validaciones específicas para el campo de email (formatos válidos e inválidos).')
b('validation_exporter.py: exporta los resultados de validación a un archivo Excel formateado, ordenando los errores primero y aplicando colores para facilitar la revisión.')
p('Las reglas de validación por campo se almacenan en archivos JSON separados por país en la carpeta json/: field_validation_rules_Argentina.json, field_validation_rules_Brasil.json, etc.')

# ── 13 ────────────────────────────────────────────────────────────────────────
h2('13. Interfaz gráfica de validación de campos — pestaña QA')
p('Archivo: interface/field_validation_ui.py')
p('El módulo de validación del punto anterior se opera desde una pestaña dedicada dentro de la interfaz gráfica, sin necesidad de usar la línea de comandos ni editar archivos de configuración manualmente.')
p('Funcionalidades de la pestaña QA:')
b('Selector de país, URL del formulario y navegador para configurar la sesión de validación.')
b('Botón para iniciar el descubrimiento de campos: el sistema abre el formulario en Selenium, lo recorre paso a paso y lista todos los campos encontrados (ID HTML, tipo de campo, label, opciones si es dropdown, si es requerido, etc.).')
b('Para cada campo descubierto, muestra las reglas de validación configuradas y el resultado de cada caso de prueba (si pasó o falló, qué mensaje de error apareció en el formulario vs cuál se esperaba).')
b('Permite editar y guardar las reglas de validación de cada campo directamente desde la UI, exportándolas automáticamente al JSON del país correspondiente para futuras ejecuciones.')
b('Genera reportes de validación en formato Excel con colores exportables para revisión por parte del equipo de QA o del cliente.')
b('Permite correr la validación completa de todos los campos o enfocarse en un campo específico para agilizar el debugging.')

# ── 14 ────────────────────────────────────────────────────────────────────────
h2('14. Consola de logs en tiempo real integrada en la interfaz')
p('Archivo: interface/console_widget.py')
p('La interfaz incluye un panel de consola que muestra todos los mensajes del sistema en tiempo real, sin necesidad de tener una ventana de terminal adicional abierta.')
p('Características técnicas:')
b('Redirige sys.stdout mediante un objeto _TeeStream que captura todos los print() del motor de automatización sin modificarlos.')
b('Thread-safe: dado que Selenium corre en un hilo separado al de la UI, las actualizaciones del widget se encolan en el hilo principal de Tkinter usando root.after(0, ...), evitando errores de concurrencia o congelamiento de la interfaz.')
b('Incluye un botón "Limpiar" para borrar el historial visible en pantalla sin afectar el log en disco ni las ejecuciones en curso.')
b('El stdout original se preserva: los mensajes también aparecen en la terminal del sistema para desarrolladores que corren la app desde la línea de comandos.')
p('Esto permite que el operador pueda ver exactamente qué está haciendo el sistema en cada momento (qué campo está llenando, qué URL navegó, si hubo un error y cuál fue) sin necesidad de salir de la interfaz gráfica.')

# ── 15 ────────────────────────────────────────────────────────────────────────
h2('15. Modo headless — ejecución sin ventana visible del navegador')
p('Archivo: core/browser_manager.py')
p('El modo headless permite que el navegador corra completamente en segundo plano, sin abrir ninguna ventana en pantalla. Esto es útil para las ejecuciones autónomas programadas (cuando no hay nadie mirando el equipo) o cuando el operador quiere correr el script sin interrumpir su trabajo en el equipo.')
p('Soporte por navegador:')
b('Chrome: argumento --headless=new (sintaxis nueva, compatible con Chrome 112 en adelante).')
b('Firefox: argumento -headless.')
b('Edge: argumento --headless.')
p('Cómo activarlo: el modo headless se activa pasando headless=True al método create_browser() de browser_manager.py. Esta opción está disponible en el código pero actualmente no se expone como un control visible en la interfaz gráfica. Para habilitarlo es necesario modificar el parámetro en el código del runner correspondiente o en la configuración de la ejecución autónoma.')
p('Las capturas de pantalla siguen funcionando correctamente en modo headless: Chrome y Edge usan el método de scroll+merge, y Firefox headless usa igualmente su método nativo de captura full-page.')

# ── 16 ────────────────────────────────────────────────────────────────────────
h2('16. Auto-descarga de drivers de navegador como fallback')
p('Archivo: core/browser_manager.py')
p('El sistema incluye en la carpeta /drivers/ los binarios de ChromeDriver, GeckoDriver (Firefox) y EdgeDriver necesarios para que Selenium controle el navegador. Sin embargo, si la versión del driver incluido no coincide con la del navegador instalado en el equipo (algo que ocurre cuando el navegador se actualiza automáticamente), Selenium falla con un error de "version mismatch".')
p('La mejora implementada:')
b('Si el driver local falla al iniciar, el sistema detecta el tipo de error.')
b('Como fallback, intenta usar la librería webdriver-manager (incluida como dependencia) para descargar automáticamente la versión correcta del driver que corresponde al navegador instalado en ese momento.')
b('Si el fallback también falla (por ejemplo, sin conexión a internet), muestra un mensaje de error claro y comprensible con instrucciones sobre qué hacer.')
b('Excepción importante: en el ejecutable .exe compilado con PyInstaller (modo frozen), solo se usan los drivers incluidos en el paquete. No se intenta la descarga automática, porque el entorno frozen puede no tener permisos de escritura en la carpeta temporal o acceso a internet en todos los casos de uso.')
p('Esta mejora reduce significativamente los problemas de configuración del entorno para nuevos usuarios o después de actualizaciones de navegador.')

# ── 17 ────────────────────────────────────────────────────────────────────────
h2('17. Executor autónomo mejorado con programación y notificaciones')
p('Archivos: autonomous_runner.py y utils/scheduling.py')
p('El executor autónomo permite programar ejecuciones desatendidas: el sistema corre automáticamente a la fecha y hora configuradas sin que el operador tenga que abrir la aplicación ni estar presente.')
p('Mejoras incorporadas:')
b('Lock global con Windows Mutex (Global\\OsocioFormAutomationAutonomous): evita que se inicien múltiples instancias del executor en paralelo accidentalmente. Si ya hay una instancia ejecutando y el sistema intenta lanzar otra (por ejemplo, por un doble clic), la segunda detecta el mutex y termina limpiamente sin duplicar trabajo.')
b('Email automático de resultados consolidados: al finalizar la ejecución de todos los países programados, envía automáticamente un email con el resumen consolidado de todos los países. El destinatario y la configuración SMTP se toman de json/config_global.json.')
b('Integración con scheduling.py: carga la programación persistida (países, navegadores, viewports, fecha y hora objetivo), la ejecuta y luego la elimina para evitar que se repita la misma ejecución en la siguiente revisión del scheduler.')
b('Manejo seguro de Unicode en consola (_safe_console_print): en Windows, la consola puede no soportar ciertos caracteres especiales (tildes, eñes) dependiendo del encoding del sistema. Esta función sanitiza los mensajes antes de imprimirlos para evitar errores de UnicodeEncodeError que interrumpirían el proceso desatendido.')
b('Ejecución como proceso hijo desacoplado: el executor se lanza como un subprocess separado desde la UI (con CREATE_NO_WINDOW en Windows), de modo que si el operador cierra la interfaz gráfica, el executor sigue corriendo en segundo plano hasta terminar.')

# ── 18 ────────────────────────────────────────────────────────────────────────
h2('18. Resumen consolidado de resultados multi-país')
p('Archivo: autonomous_runner.py — integrado en el flujo de ejecución autónoma')
p('Cuando el executor autónomo termina de procesar todos los países de la programación, genera automáticamente un archivo resumen_consolidado.xlsx en la carpeta resultados/.')
p('Este Excel consolidado:')
b('Agrega en una sola hoja los resultados de todos los países ejecutados en esa corrida, sin importar cuántos países hayan corrido.')
b('Incluye columnas como: país, URL del formulario, resultado (Éxito / Error), mensaje de error si hubo alguno, y rutas a los archivos de capturas de pantalla.')
b('Permite ver de un solo vistazo cuántos leads se enviaron exitosamente en cada país y cuáles fallaron, sin necesidad de abrir 9 archivos Excel separados.')
b('Este archivo consolidado es el que se adjunta en el email de resultados que el executor envía al finalizar la corrida.')

# ── 19 ────────────────────────────────────────────────────────────────────────
h2('19. Retry automático por error transitorio de envío')
p('Archivo: core/base_form_filler.py — bloque de procesamiento de filas del loop principal')
p('A veces el servidor del formulario devuelve un error temporal después de enviar: "Lo siento, ocurrió un inconveniente al realizar el envío del formulario en este momento. Por favor, inténtalo nuevamente." Este mensaje indica un problema temporal del servidor, no un error en los datos del lead.')
p('El sistema detecta y maneja este caso automáticamente:')
b('Después de cada envío, verifica si el texto de error transitorio aparece en la página de respuesta.')
b('Si lo detecta, recarga la landing page completa, espera a que el iframe del formulario esté listo, y rellena todos los campos nuevamente con los mismos datos de la misma fila del Excel.')
b('Límite de intentos: 2 en total (el intento original más 1 retry). Si el segundo intento también falla con el mismo error, el resultado se registra como error definitivo en el Excel, indicando el mensaje de error recibido.')
b('El texto de detección es "ocurrió un inconveniente al realizar el envío", lo suficientemente específico para no confundirse con otros mensajes de validación o error de datos.')
p('Esta mejora reduce la tasa de falsos negativos: leads que son perfectamente válidos pero que fallaron por un problema transitorio del servidor quedan con una segunda oportunidad de ser enviados.')

# ── 20 ────────────────────────────────────────────────────────────────────────
h2('20. Módulos utilitarios nuevos')
p('Archivos: utils/paths.py, utils/scheduling.py, utils/migrate_validation_rules.py, utils/popup_logger.py, utils/fixed_field_mapping_store.py')

p('utils/paths.py — Resolución de rutas compatible con ejecución normal y PyInstaller:')
b('Cuando la aplicación se compila con PyInstaller en un .exe, los archivos quedan empaquetados en una carpeta temporal (sys._MEIPASS). Las rutas relativas normales dejan de funcionar porque el ejecutable corre desde una ubicación diferente.')
b('Este módulo expone constantes (BASE_DIR, DATA_DIR, RESULTS_DIR, JSON_DIR, DRIVERS_DIR) que resuelven correctamente la ruta tanto en modo desarrollo (python run.py) como en modo compilado (.exe), detectando automáticamente el contexto de ejecución.')
b('Todos los módulos del sistema importan estas constantes en lugar de construir rutas hardcodeadas, garantizando que el .exe funcione correctamente sin modificaciones.')

p('utils/scheduling.py — Gestión de programaciones de ejecución automática:')
b('guardar_programacion(dict): persiste la programación futura en un archivo JSON. Guarda: lista de países, navegadores, viewports, y la fecha y hora objetivo de ejecución.')
b('cargar_programacion(): recupera la programación guardada. Si no existe ninguna, retorna None.')
b('limpiar_programacion(): elimina el archivo de programación una vez ejecutada, para evitar que la misma corrida se repita la próxima vez que el scheduler revise.')

p('utils/migrate_validation_rules.py — Migración de reglas de validación:')
b('Herramienta de mantenimiento que convierte reglas de validación del formato antiguo (un único JSON global para todos los países) al formato nuevo (un JSON separado por país).')

p('utils/popup_logger.py — Sistema de logging centralizado:')
b('Gestiona el archivo de log en disco, resolviendo la ruta correctamente tanto en modo desarrollo como en modo .exe.')
b('Provee notificaciones de popup para errores críticos que el operador debe atender.')

p('utils/fixed_field_mapping_store.py — Persistencia de mappings de campos sin editar código:')
b('load_effective_country_form_config(country_name): combina la configuración base de un país (definida en country_configs.py) con los overrides guardados por el operador en json/fixed_field_mappings.json. El resultado es el mapping efectivo que usa el motor en runtime.')
b('Permite cambiar los IDs HTML de los campos del formulario (cuando el sitio de GM los actualiza) o agregar campos nuevos, directamente desde la UI, y que esos cambios persistan entre sesiones sin necesidad de editar el código fuente ni reiniciar la aplicación.')
b('build_excel_columns_for_country(country_name): genera la lista de columnas que debe tener el Excel de entrada para un país, alineada con el mapping efectivo actual.')

# ── 21 ────────────────────────────────────────────────────────────────────────
h2('21. Mejoras en el proceso de build y distribución del ejecutable')
p('Archivos: build.bat y FormAutomation.spec')
p('El script de compilación que genera el ejecutable .exe fue mejorado para ser más robusto y reproducible en distintos entornos de desarrollo:')
b('Detección dinámica de Python instalado: el script prueba automáticamente versiones de Python desde 3.9 hasta 3.14, y también busca en el PATH del sistema. Ignora los stubs de Microsoft Store (WindowsApps) que son launchers vacíos que confunden a otros scripts de build.')
b('Verificación de errores en cada paso: después de cada comando pip install, verifica el errorlevel y detiene el proceso inmediatamente si hubo un fallo, mostrando un mensaje de error claro. Esto evita que la compilación continúe con dependencias faltantes o mal instaladas y produzca un .exe roto.')
b('Generación automática del ZIP portable: al finalizar la compilación con PyInstaller, el script genera automáticamente un archivo .zip que contiene el ejecutable y todos los archivos necesarios usando el comando PowerShell Compress-Archive. Esto facilita la distribución del programa a nuevos usuarios.')
b('Salida limpia al finalizar: muestra las rutas exactas del .exe generado, la carpeta portable y el archivo .zip, para que el desarrollador sepa exactamente qué archivos usar y dónde están.')
p('El archivo FormAutomation.spec configura PyInstaller para incluir correctamente en el paquete: los tres drivers de navegador (chromedriver.exe, geckodriver.exe, msedgedriver.exe), los archivos de configuración JSON, los Excel de datos iniciales vacíos y todos los recursos gráficos (iconos, imágenes de la interfaz).')

# ══════════════════════════════════════════════════════════════════════════════
# PUNTO DE PARTIDA PARA ONBOARDING
# ══════════════════════════════════════════════════════════════════════════════
h2('Punto de partida recomendado para onboarding')
p('Si un desarrollador nuevo tiene poco tiempo, el orden recomendado para leer el proyecto es el siguiente. La arquitectura fue refactorizada: ya no existen archivos individuales por país (Formulario_*_Base.py ni Formulario_*_Main.py). Todo está centralizado en módulos genéricos.')

h3('1. run.py')
p('Archivo: run.py')
b('Qué hace: punto de entrada principal. Parsea argumentos de línea de comandos (--autonomous, --run-country, --environment, --no-email) y decide si lanzar la UI, ejecutar un país directamente o iniciar el executor autónomo.')
b('Inputs: argumentos de consola opcionales. Sin argumentos, abre la interfaz gráfica.')
b('Outputs: lanza interfaz gráfica, ejecución de un país puntual, o el scheduler autónomo.')
b('Por qué es importante: es la puerta de entrada a todo el sistema, tanto en modo desarrollo como en el .exe compilado.')

h3('2. interface/main_interface.py')
p('Archivo: interface/main_interface.py')
b('Qué hace: construye la aplicación Tkinter con todas sus pestañas (Formularios, Generar Excels, Resultados, Ejecución Autónoma, Validación de Campos). Expone los controles para seleccionar país, navegador, viewport y lanzar ejecuciones.')
b('Inputs: no recibe parámetros directos; lee configuraciones de json/ y data/ al iniciar.')
b('Outputs: abre la UI y deja registrado el flujo operativo del usuario. Al ejecutar un país, lanza un hilo con el runner correspondiente.')
b('Por qué es importante: es la puerta de entrada humana del sistema. Sin esta función no existe operación diaria del proyecto.')

h3('3. forms/_runner_common.py')
p('Archivo: forms/_runner_common.py')
b('Qué hace: fábrica de runners por país. La función get_runner(country_name) genera dinámicamente la función run_formularios_<País>() para cualquier país, creando una subclase anónima de GenericCountryBase. Reemplaza los 9 archivos Formulario_*_Main.py que existían en versiones anteriores.')
b('Inputs: nombre del país (string).')
b('Outputs: función callable con firma (browser, viewport, headless, enviar_email) que ejecuta el formulario del país.')
b('Por qué es importante: es el puente entre la selección de país en la UI y la ejecución concreta. Entender get_runner() es entender cómo el sistema pasa de "el usuario clickeó Argentina" a "se abre Chrome y se llena el formulario".')

h3('4. core/country_configs.py')
p('Archivo: core/country_configs.py')
b('Qué hace: define el diccionario COUNTRY_CONFIGS con la configuración declarativa de los 9 países. Cada entrada contiene el archivo Excel de entrada, el field_mapping (campos del formulario con su tipo, ID HTML y columna del Excel), los required_fields y flags especiales por país.')
b('Inputs: no recibe parámetros; es un módulo de configuración estático.')
b('Outputs: función get_country_config(country_name) que retorna el dict de configuración del país.')
b('Por qué es importante: reemplaza los 9 archivos Formulario_*_Base.py. Aquí se define qué campos existen en cada formulario y cómo se leen del Excel. Es el primer lugar a revisar cuando un campo falla o necesita ajuste.')

h3('5. core/generic_country_base.py')
p('Archivo: core/generic_country_base.py')
b('Qué hace: clase GenericCountryBase que hereda de BaseFormFiller. Carga la configuración de cualquier país desde country_configs.py y la pasa al motor base. Reemplaza los 9 archivos Formulario_*_Base.py individuales.')
b('Inputs: country_name (string), browser, viewport, headless.')
b('Outputs: instancia lista para llamar a .run() y ejecutar el formulario del país.')
b('Por qué es importante: es la glue entre la configuración declarativa (country_configs.py) y el motor de ejecución (base_form_filler.py). Leerlo deja claro que toda la lógica real vive en BaseFormFiller.')

h3('6. core/base_form_filler.py')
p('Archivo: core/base_form_filler.py')
b('Qué hace: motor central de la automatización. Contiene run(), process_landing_page(), find_and_position_to_form(), fill_form_fields_auto_step(), safe_select_option_if_visible(), submit_and_verify_form(), write_tracked_fields_to_sheet(), _generate_brazil_document() y toda la lógica de Selenium: apertura de landing, manejo de iframes, cookies, pasos del formulario, dependencias de dropdowns, capturas y guardado de resultados.')
b('Inputs: recibe el dict de configuración efectiva del país (config) al instanciarse.')
b('Outputs: Excel de resultados, screenshots, logs en consola y cierre ordenado del navegador.')
b('Por qué es importante: es el archivo más importante del proyecto. Todo el comportamiento real del llenado, envío y captura de evidencias vive aquí. Cualquier bug de automatización se diagnostica en este archivo.')

h3('7. utils/fixed_field_mapping_store.py')
p('Archivo: utils/fixed_field_mapping_store.py')
b('Qué hace: centraliza la configuración efectiva por país combinando los defaults de country_configs.py con los overrides persistidos en json/fixed_field_mappings.json. Expone load_effective_country_form_config(country_name) y build_excel_columns_for_country(country_name).')
b('Inputs: nombre del país.')
b('Outputs: dict de configuración efectiva con field_mapping, required_fields, excel_file, etc.')
b('Por qué es importante: desacopla la operación del código fuente. Permite cambiar IDs de campos o agregar nuevos desde la UI sin editar country_configs.py. Es la pieza que hace que los cambios operativos no requieran commits.')

h3('8. interface/helpers_interface.py')
p('Archivo: interface/helpers_interface.py')
b('Qué hace: funciones auxiliares de la UI. Incluye creación y sincronización de Excels (abrir_excel), envío de email con resultados (enviar_email_resultados, enviar_email_resultados_consolidados), configuración SMTP, gestión de IDs dinámicos desde la UI y consolidación de resultados multi-país.')
b('Inputs: rutas de archivos, configuraciones de email, datos de resultados según la función.')
b('Outputs: True/False de éxito en operaciones de archivo o email; archivos Excel creados o actualizados.')
b('Por qué es importante: cierra el ciclo operativo del sistema. Sin este módulo no habría creación automática de Excels ni notificaciones de resultados.')

h3('9. validation/selenium_validation_runner.py')
p('Archivo: validation/selenium_validation_runner.py')
b('Qué hace: motor de QA automatizado. Abre formularios reales con Selenium, descubre todos los campos visibles (incluyendo pasos intermedios), ejecuta casos de prueba (campo vacío, valores inválidos, bordes), captura los mensajes de error del formulario y compara con los esperados según json/field_validation_rules_*.json.')
b('Inputs: field_mapping, URLs del formulario, navegador, viewport y parámetros de ejecución.')
b('Outputs: estructura con rows (resultados por campo), fields (campos descubiertos), error_rows (fallos), unmapped_ids y summary.')
b('Por qué es importante: es el motor del subsistema de QA. Permite verificar que los formularios validan correctamente sin intervención manual.')

h3('10. autonomous_runner.py')
p('Archivo: autonomous_runner.py')
b('Qué hace: loop residente que carga la programación desde utils/scheduling.py, ejecuta todos los países/navegadores/viewports programados usando get_runner() de forms/_runner_common.py, genera el resumen consolidado y envía el email de resultados al finalizar. Usa un Windows Mutex para evitar instancias duplicadas.')
b('Inputs: no recibe parámetros directos; lee json/ para la programación guardada por scheduling.py.')
b('Outputs: ejecuciones de formularios, resultados/resumen_consolidado.xlsx y email de resultados.')
b('Por qué es importante: es el corazón del modo desatendido. Permite que el sistema corra solo a una hora programada sin que nadie esté presente.')

p('Con esa secuencia se entiende casi todo el sistema sin necesidad de recorrer cada archivo individualmente.')

# ══════════════════════════════════════════════════════════════════════════════
# OBSERVACIONES ÚTILES
# ══════════════════════════════════════════════════════════════════════════════
h2('Observaciones útiles para un desarrollador nuevo')
b('La UI es importante para operar, pero no contiene la lógica central de automatización. El comportamiento real del llenado se decide en core/base_form_filler.py.')
b('Los archivos forms/Formulario_*_Main.py ya no existen como archivos individuales. La función get_runner(country_name) en forms/_runner_common.py los reemplaza y genera runners dinámicamente para cualquier país.')
b('Los archivos core/Formulario_*_Base.py tampoco existen. Fueron reemplazados por core/generic_country_base.py (clase única) + core/country_configs.py (configuración declarativa).')
b('Los cambios operativos de mapping (cuando un ID de campo del formulario cambia en producción) no requieren editar código. Se persisten en json/fixed_field_mappings.json desde la UI y utils/fixed_field_mapping_store.py los aplica en runtime.')
b('Los IDs dinámicos, valores fijos para campos y dependencias configurables por el operador viven en json/ids_dinamicos.json. No todo está hardcodeado en el código.')
b('La selección de dropdowns es siempre exacta: el valor del Excel debe coincidir textualmente con la opción visible del formulario. Si no coincide, el campo falla y se reporta en consola y en el Excel de resultados (columna "Resultado"). No hay matching aproximado.')
b('El módulo de validación (validation/) es un subsistema separado útil para QA y descubrimiento de problemas del DOM. No interfiere con el flujo de carga de leads.')
b('El scheduler autónomo no reemplaza la UI: la complementa para corridas programadas. La UI sigue siendo el método principal de operación manual.')
b('json/config_global.json centraliza los timeouts del sistema (page_load, element_wait, form_submit, step_change, dependency_dropdown). Ajustar estos valores afecta el comportamiento de todos los países.')
b('Para agregar un nuevo país: agregar una entrada en COUNTRY_CONFIGS (core/country_configs.py) y crear el Excel de entrada en data/. No se necesita ningún otro archivo nuevo.')

# ══════════════════════════════════════════════════════════════════════════════
# FLUJO RESUMIDO
# ══════════════════════════════════════════════════════════════════════════════
h2('Flujo resumido con funciones')
p('Este es el recorrido principal del código en el orden en que un desarrollador nuevo debería entenderlo. La arquitectura fue refactorizada: ya no hay archivos individuales por país.')

h3('1. run.py')
p('Arranca la aplicación. Sin argumentos de CLI, delega en iniciar_interfaz(). Con --autonomous, lanza el executor autónomo. Con --run-country, ejecuta un país puntual sin abrir la UI.')

h3('2. interface.main_interface.iniciar_interfaz()')
p('Levanta la UI Tkinter con todas sus pestañas. Expone controles para seleccionar país, navegador y viewport. Al presionar "Ejecutar", llama a ejecutar_script_configurable().')

h3('3. interface.main_interface.ejecutar_script_configurable()')
p('Traduce la acción del usuario (país + navegador + viewport elegidos en UI) en una ejecución concreta. Obtiene el runner del país llamando a forms._runner_common.get_runner(country_name) y lo ejecuta en un hilo separado.')

h3('4. forms._runner_common.get_runner(country_name)')
p('Fábrica de runners. Crea dinámicamente una subclase de GenericCountryBase para el país solicitado y devuelve una función callable con firma (browser, viewport, headless, enviar_email). No existe un archivo Formulario_Argentina_Main.py — este módulo lo reemplaza para los 9 países.')

h3('5. utils.fixed_field_mapping_store.load_effective_country_form_config()')
p('Resuelve el mapping que realmente se usará en runtime. Combina la configuración base de country_configs.py con los overrides guardados por el operador en json/fixed_field_mappings.json. El resultado es el dict de configuración efectiva que recibe BaseFormFiller.')

h3('6. core.base_form_filler.run()')
p('Orquestador principal del motor Selenium. Crea el navegador (browser_manager), abre el Excel de entrada, recorre fila por fila y coordina el procesamiento end-to-end. Al finalizar todas las filas, cierra el navegador y guarda el Excel de resultados.')

h3('7. core.base_form_filler.process_landing_page()')
p('Para cada fila: navega a la landing page, espera su carga completa, detecta y cierra banners de cookies/popups, y toma la captura inicial de evidencia.')

h3('8. core.base_form_filler.find_and_position_to_form()')
p('Localiza el iframe correcto del formulario dentro de la landing (por URL esperada del iframe), hace scroll hasta él y cambia el contexto de Selenium al interior del iframe. Si no lo encuentra, la fila se marca como error.')

h3('9. core.base_form_filler.capture_error_messages()')
p('Dispara el envío del formulario vacío para capturar los mensajes de validación iniciales. Guarda el screenshot como evidencia del estado base del formulario antes de llenarlo.')

h3('10. core.base_form_filler.fill_form_fields_auto_step()')
p('Ejecuta el llenado real del formulario. Detecta si es de uno o múltiples pasos buscando botones "Siguiente/Seguinte/Next/Continuar". Llena los campos visibles del paso actual, avanza al siguiente y repite hasta el paso final (máx. 15 iteraciones).')

h3('11. core.base_form_filler._fill_visible_fields_from_mapping()')
p('Decide qué campo completar en el step actual. Itera el field_mapping, verifica visibilidad, resuelve dependencias (región→ciudad→concesionario), usa safe_select_option_if_visible() para selects y send_keys / JS injection para textos. Solo toca campos visibles en el DOM actual.')

h3('12. core.base_form_filler._auto_fill_unmapped_dropdowns()')
p('Absorbe selects e inputs visibles que no estaban en el mapping. Si hay un ID dinámico configurado para ese campo en ids_dinamicos.json, lo usa. Si no, selecciona una opción válida aleatoria. Hace el sistema resiliente ante campos nuevos no mapeados.')

h3('13. core.base_form_filler.submit_and_verify_form()')
p('Envía el formulario, espera a que aparezca el div#thank-you con display:block (TY page). Si en 15 segundos no aparece, verifica errores visibles o campos todavía presentes para determinar el resultado. Si detecta el mensaje de error transitorio del servidor, activa el retry automático.')

h3('14. core.base_form_filler.write_tracked_fields_to_sheet()')
p('Escribe en el Excel de resultados los valores efectivamente usados para cada campo (mapeados y dinámicos). Los campos no mapeados generan columnas nuevas con formato ID::{field_id}. Deja trazabilidad completa de qué se ingresó en cada fila.')

h3('15. interface.helpers_interface.enviar_email_resultados()')
p('Si el email está habilitado en la configuración: analiza el Excel generado, comprime las capturas de pantalla en un ZIP y envía el paquete de evidencias por Outlook al destinatario configurado. Cierra el ciclo operativo de cada ejecución.')

# ── Cierre ────────────────────────────────────────────────────────────────────
p('')
p('Fecha de actualización: Junio 2026')

doc.save('Arquitectura_Consolidada_ Osocio Form Automation_FINAL3.docx')
print('Guardado como FINAL3.docx')
